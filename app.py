"""
Generador de Informe de Citas — BCCT / UNAM
Convierte XMLs de EndNote en informes de citas (UNAM o SNI/SECIHTI)
y automatiza el registro en la plataforma Rizoma con Playwright.
Ejecutar: python app.py  →  http://localhost:5000
"""

import json
import os
import re
import glob
import asyncio
import threading
from pathlib import Path
from datetime import datetime, date
from xml.etree import ElementTree as ET
from flask import Flask, render_template_string, jsonify, request, send_from_directory

# ── Intentar importar python-docx y reportlab (opcionales para exportación) ──
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False
    print("⚠️ python-docx no disponible. Exportación a DOCX deshabilitada.")

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import cm
    PDF_OK = True
except ImportError:
    PDF_OK = False
    print("⚠️ reportlab no disponible. Exportación a PDF deshabilitada.")

# ── Intentar importar Playwright (opcional para automatización) ──
try:
    from playwright.async_api import async_playwright
    PLAYWRIGHT_OK = True
except ImportError:
    PLAYWRIGHT_OK = False
    print("⚠️ Playwright no disponible. Automatización de Rizoma deshabilitada.")

app = Flask(__name__)
UPLOAD_DIR = Path("uploads")
OUTPUT_DIR = Path("outputs")
PROFILES_FILE = Path("profiles.json")
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

# ── Estado global para la automatización Rizoma ──
_task = {"active": False, "log": [], "status": "idle", "progress": 0, "total": 0}

# ══════════════════════════════════════════════════════════════════════════════
# MODELOS DE DATOS
# ══════════════════════════════════════════════════════════════════════════════

def load_profiles() -> dict:
    if PROFILES_FILE.exists():
        return json.loads(PROFILES_FILE.read_text(encoding="utf-8"))
    return {}

def save_profiles(profiles: dict):
    PROFILES_FILE.write_text(json.dumps(profiles, indent=2, ensure_ascii=False), encoding="utf-8")

def get_profile(name: str) -> dict | None:
    return load_profiles().get(name)

# ══════════════════════════════════════════════════════════════════════════════
# PARSER DE XML ENDNOTE
# ══════════════════════════════════════════════════════════════════════════════

def _txt(elem, path: str) -> str:
    """Extrae texto limpio de un sub-elemento XML de EndNote."""
    node = elem.find(path)
    if node is None:
        return ""
    parts = []
    for child in node.iter():
        if child.text:
            parts.append(child.text.strip())
        if child.tail:
            parts.append(child.tail.strip())
    return " ".join(p for p in parts if p)

def parse_endnote_xml(xml_path: str) -> list[dict]:
    """
    Parsea un XML exportado de EndNote y devuelve lista de citas.
    """
    try:
        tree = ET.parse(xml_path)
        root = tree.getroot()
    except ET.ParseError:
        return []

    citas = []
    for rec in root.findall(".//record"):
        authors = []
        for a in rec.findall(".//authors/author"):
            name = _txt(a, "style") or "".join(a.itertext()).strip()
            if name:
                authors.append(name)

        title = _txt(rec, "titles/title") or _txt(rec, "titles/title/style")
        journal = _txt(rec, "titles/secondary-title") or _txt(rec, "periodical/full-title")
        year = _txt(rec, "dates/year")
        volume = _txt(rec, "volume")
        issue = _txt(rec, "number")
        pages = _txt(rec, "pages")
        doi = _txt(rec, "electronic-resource-num")
        isbn = _txt(rec, "isbn")
        database = _txt(rec, "remote-database-name")

        url = ""
        for url_elem in rec.findall(".//urls/related-urls/url"):
            url = "".join(url_elem.itertext()).strip()
            if url:
                break

        notes_raw = _txt(rec, "research-notes")
        notes_upper = notes_raw.upper().strip()
        if notes_upper == "B":
            tipo = "B"
        elif notes_upper == "C":
            tipo = "C"
        else:
            tipo = "A"

        autor_str = "; ".join(authors[:6])
        if len(authors) > 6:
            autor_str += " et al."
        apa = f"{autor_str} ({year}). {title}. {journal}"
        if volume:
            apa += f", {volume}"
            if issue:
                apa += f"({issue})"
        if pages:
            apa += f", {pages}"
        if doi:
            apa += f". https://doi.org/{doi.lstrip('https://doi.org/')}"

        citas.append({
            "authors": authors,
            "title": title,
            "journal": journal,
            "year": year,
            "volume": volume,
            "issue": issue,
            "pages": pages,
            "doi": doi,
            "isbn": isbn,
            "url": url,
            "database": database,
            "tipo": tipo,
            "apa": apa.strip(),
        })
    return citas

# ══════════════════════════════════════════════════════════════════════════════
# GENERADOR DE INFORME
# ══════════════════════════════════════════════════════════════════════════════

def generar_informe(profile: dict, modalidad: str) -> dict:
    publicaciones = profile.get("publicaciones", [])
    total_a = total_b = total_c = 0
    docs_procesados = []

    for pub in publicaciones:
        citas_raw = pub.get("citas", [])

        if modalidad == "SNI":
            citas = [c for c in citas_raw if c["tipo"] in ("A", "B")]
        else:
            citas = citas_raw

        cnt_a = sum(1 for c in citas if c["tipo"] == "A")
        cnt_b = sum(1 for c in citas if c["tipo"] == "B")
        cnt_c = sum(1 for c in citas if c["tipo"] == "C")

        total_a += cnt_a
        total_b += cnt_b
        total_c += cnt_c

        urls_pub = pub.get("urls_citas", {}) if modalidad == "SNI" else {}

        docs_procesados.append({
            "referencia": pub.get("referencia_apa", ""),
            "doi": pub.get("doi", ""),
            "fi": pub.get("fi", ""),
            "quartil": pub.get("quartil", ""),
            "issn_e": pub.get("issn_e", ""),
            "issn_i": pub.get("issn_i", ""),
            "isbn": pub.get("isbn", ""),
            "cnt_a": cnt_a,
            "cnt_b": cnt_b,
            "cnt_c": cnt_c,
            "total_citas": len(citas),
            "urls_citas": urls_pub,
            "citas": citas,
        })

    return {
        "modalidad": modalidad,
        "investigador": profile.get("nombre", ""),
        "entidad": profile.get("entidad", ""),
        "especialista": profile.get("especialista", ""),
        "revisor": profile.get("revisor", ""),
        "fecha_revision": profile.get("fecha_revision", str(date.today())),
        "fuentes": profile.get("fuentes", []),
        "total_a": total_a,
        "total_b": total_b,
        "total_c": total_c,
        "total_general": total_a + total_b + total_c,
        "documentos": docs_procesados,
    }

# ══════════════════════════════════════════════════════════════════════════════
# EXPORTADOR DOCX
# ══════════════════════════════════════════════════════════════════════════════

def exportar_docx(informe: dict, output_path: str):
    if not DOCX_OK:
        raise Exception("python-docx no está instalado")
    
    doc = Document()
    
    # Título
    p = doc.add_paragraph()
    run = p.add_run(f"Informe de Citas - {informe['modalidad']}")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(97, 18, 50)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Investigador: {informe['investigador']}")
    doc.add_paragraph(f"Entidad: {informe['entidad']}")
    doc.add_paragraph(f"Fecha: {informe['fecha_revision']}")
    doc.add_paragraph()
    
    # Totales
    doc.add_heading("Resumen de Citas", level=2)
    doc.add_paragraph(f"Total Tipo A (Individual): {informe['total_a']}")
    doc.add_paragraph(f"Total Tipo B (Colaborativa): {informe['total_b']}")
    if informe['modalidad'] == 'UNAM':
        doc.add_paragraph(f"Total Tipo C (Autocita): {informe['total_c']}")
    doc.add_paragraph(f"TOTAL GENERAL: {informe['total_general']}")
    doc.add_paragraph()
    
    # Documentos
    doc.add_heading("Documentos Procesados", level=2)
    for idx, doc_info in enumerate(informe['documentos'], 1):
        doc.add_heading(f"Documento {idx}", level=3)
        doc.add_paragraph(f"Referencia: {doc_info['referencia']}")
        doc.add_paragraph(f"Citas tipo A: {doc_info['cnt_a']}")
        doc.add_paragraph(f"Citas tipo B: {doc_info['cnt_b']}")
        if informe['modalidad'] == 'UNAM':
            doc.add_paragraph(f"Citas tipo C: {doc_info['cnt_c']}")
        doc.add_paragraph()
    
    doc.save(output_path)

# ══════════════════════════════════════════════════════════════════════════════
# AUTOMATIZACIÓN RIZOMA
# ══════════════════════════════════════════════════════════════════════════════

async def automatizar_rizoma(profile: dict, informe: dict, credenciales: dict):
    """
    Función placeholder para automatización de Rizoma.
    Requiere playwright y rizoma_automation.py
    """
    if not PLAYWRIGHT_OK:
        _task["log"].append("❌ Playwright no está instalado. Ejecuta: pip install playwright")
        _task["log"].append("❌ Luego ejecuta: python -m playwright install chromium")
        _task["status"] = "error"
        return
    
    try:
        # Importar el módulo de automatización
        from rizoma_automation import automatizar_rizoma_completo
        await automatizar_rizoma_completo(profile, credenciales, _task)
    except ImportError:
        _task["log"].append("❌ No se encontró rizoma_automation.py")
        _task["log"].append("❌ Asegúrate de tener el archivo en la misma carpeta que app.py")
        _task["status"] = "error"
    except Exception as e:
        _task["log"].append(f"❌ Error en automatización: {str(e)}")
        _task["status"] = "error"

# ══════════════════════════════════════════════════════════════════════════════
# PLANTILLA HTML
# ══════════════════════════════════════════════════════════════════════════════

HTML = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Sistema Rizoma - BCCT/UNAM</title>
<style>
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: 'Segoe UI', Arial, sans-serif; background: #f8f9fa; color: #333; }
.container { max-width: 1400px; margin: 0 auto; padding: 20px; }
header { background: linear-gradient(135deg, #611232 0%, #8B1538 100%); color: white; padding: 30px 20px; text-align: center; margin-bottom: 30px; border-radius: 10px; }
header h1 { font-size: 2em; margin-bottom: 10px; }
.tabs { display: flex; gap: 10px; margin-bottom: 20px; border-bottom: 2px solid #ddd; }
.tab { padding: 12px 24px; cursor: pointer; background: white; border: none; border-bottom: 3px solid transparent; font-size: 1em; transition: all 0.3s; }
.tab.active { border-bottom-color: #611232; color: #611232; font-weight: bold; }
.tab:hover { background: #f8f9fa; }
.tab-content { display: none; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
.tab-content.active { display: block; }
.btn { padding: 10px 20px; background: #611232; color: white; border: none; border-radius: 5px; cursor: pointer; font-size: 1em; transition: all 0.3s; }
.btn:hover { background: #8B1538; transform: translateY(-2px); box-shadow: 0 4px 8px rgba(0,0,0,0.2); }
.btn:disabled { background: #ccc; cursor: not-allowed; transform: none; }
.btn-secondary { background: #6c757d; }
.btn-secondary:hover { background: #5a6268; }
.form-group { margin-bottom: 20px; }
.form-group label { display: block; margin-bottom: 5px; font-weight: 600; color: #555; }
.form-group input, .form-group select, .form-group textarea { width: 100%; padding: 10px; border: 1px solid #ddd; border-radius: 5px; font-size: 1em; }
.form-group textarea { min-height: 100px; resize: vertical; }
.card { background: white; padding: 20px; border-radius: 10px; box-shadow: 0 2px 5px rgba(0,0,0,0.1); margin-bottom: 20px; }
.profile-item { padding: 15px; border: 1px solid #e0e0e0; border-radius: 8px; margin-bottom: 15px; cursor: pointer; transition: all 0.3s; }
.profile-item:hover { background: #f8f9fa; border-color: #611232; transform: translateX(5px); }
.profile-item.selected { background: #fff3f6; border-color: #611232; border-width: 2px; }
.pub-item { background: #f8f9fa; padding: 15px; border-left: 4px solid #611232; margin-bottom: 15px; border-radius: 5px; }
.log-console { background: #1e1e1e; color: #d4d4d4; padding: 20px; border-radius: 8px; font-family: 'Consolas', monospace; max-height: 400px; overflow-y: auto; margin-top: 20px; }
.log-line { margin: 5px 0; padding: 5px; }
.log-ok { color: #4ec9b0; }
.log-err { color: #f48771; }
.log-info { color: #dcdcaa; }
.progress-bar { background: #e9ecef; height: 30px; border-radius: 15px; overflow: hidden; margin: 20px 0; }
.progress-fill { background: linear-gradient(90deg, #611232, #8B1538); height: 100%; transition: width 0.5s; display: flex; align-items: center; justify-content: center; color: white; font-weight: bold; }
.status-indicator { display: flex; align-items: center; gap: 10px; padding: 15px; background: #f8f9fa; border-radius: 8px; margin-bottom: 20px; }
.dot { width: 12px; height: 12px; border-radius: 50%; background: #28a745; animation: pulse 2s infinite; }
.dot.busy { background: #ffc107; }
@keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.5; } }
</style>
</head>
<body>
<div class="container">
<header>
  <h1>🎓 Sistema de Automatización Rizoma</h1>
  <p>Gestión de Citas Científicas — BCCT/UNAM</p>
</header>

<div class="tabs">
  <button class="tab active" onclick="showTab('perfiles')">📂 Perfiles</button>
  <button class="tab" onclick="showTab('publicaciones')">📚 Publicaciones</button>
  <button class="tab" onclick="showTab('xmls')">📄 Cargar XMLs</button>
  <button class="tab" onclick="showTab('informe')">📊 Generar Informe</button>
  <button class="tab" onclick="showTab('rizoma')">🤖 Automatizar Rizoma</button>
</div>

<!-- TAB: PERFILES -->
<div id="perfiles" class="tab-content active">
  <h2>Gestión de Perfiles</h2>
  <button class="btn" onclick="showNewProfileForm()">+ Crear Nuevo Perfil</button>
  
  <div id="newProfileForm" style="display:none; margin-top: 20px;">
    <div class="card">
      <h3>Nuevo Perfil de Investigador</h3>
      <div class="form-group">
        <label>Nombre Completo:</label>
        <input type="text" id="profNombre" placeholder="Ej: Dr. Juan Pérez García">
      </div>
      <div class="form-group">
        <label>Entidad/Institución:</label>
        <input type="text" id="profEntidad" placeholder="Ej: UNAM, Instituto de Investigaciones">
      </div>
      <div class="form-group">
        <label>Especialista/Evaluador:</label>
        <input type="text" id="profEspecialista" placeholder="Nombre del evaluador">
      </div>
      <div class="form-group">
        <label>Revisor:</label>
        <input type="text" id="profRevisor" placeholder="Nombre del revisor">
      </div>
      <div class="form-group">
        <label>Fuentes de Información:</label>
        <input type="text" id="profFuentes" placeholder="Ej: Scopus, Web of Science (separadas por comas)">
      </div>
      <button class="btn" onclick="crearPerfil()">Crear Perfil</button>
      <button class="btn btn-secondary" onclick="hideNewProfileForm()">Cancelar</button>
    </div>
  </div>

  <div id="profilesList" style="margin-top: 20px;"></div>
</div>

<!-- TAB: PUBLICACIONES -->
<div id="publicaciones" class="tab-content">
  <h2>Publicaciones del Perfil</h2>
  <div class="status-indicator">
    <span>Perfil Activo:</span>
    <strong id="activeName">Ninguno seleccionado</strong>
  </div>

  <button class="btn" onclick="showNewPubForm()">+ Agregar Publicación</button>

  <div id="newPubForm" style="display:none; margin-top: 20px;">
    <div class="card">
      <h3>Nueva Publicación</h3>
      <div class="form-group">
        <label>Referencia APA:</label>
        <textarea id="pubRef" placeholder="Apellido, N. (2024). Título del artículo. Nombre de la Revista, 12(3), 45-67."></textarea>
      </div>
      <div class="form-group">
        <label>DOI:</label>
        <input type="text" id="pubDoi" placeholder="10.1234/ejemplo">
      </div>
      <div class="form-group">
        <label>Factor de Impacto:</label>
        <input type="text" id="pubFi" placeholder="3.45">
      </div>
      <div class="form-group">
        <label>Cuartil:</label>
        <select id="pubQuartil">
          <option value="">Seleccionar</option>
          <option value="Q1">Q1</option>
          <option value="Q2">Q2</option>
          <option value="Q3">Q3</option>
          <option value="Q4">Q4</option>
        </select>
      </div>
      <div class="form-group">
        <label>ISSN Electrónico:</label>
        <input type="text" id="pubIssnE" placeholder="1234-5678">
      </div>
      <div class="form-group">
        <label>ISSN Impreso:</label>
        <input type="text" id="pubIssnI" placeholder="8765-4321">
      </div>
      <div class="form-group">
        <label>Nombre del archivo XML (Librería EndNote):</label>
        <input type="text" id="pubLibreria" placeholder="Publicacion_2024">
      </div>
      <button class="btn" onclick="agregarPublicacion()">Agregar</button>
      <button class="btn btn-secondary" onclick="hideNewPubForm()">Cancelar</button>
    </div>
  </div>

  <div id="pubsList" style="margin-top: 20px;"></div>
</div>

<!-- TAB: CARGAR XMLs -->
<div id="xmls" class="tab-content">
  <h2>Cargar Archivos XML de EndNote</h2>
  <div class="card">
    <p><strong>Instrucciones:</strong> Selecciona los archivos XML exportados desde EndNote. Los nombres de los archivos deben coincidir con el campo "Librería EndNote" de cada publicación.</p>
    <div class="form-group">
      <label>Seleccionar XMLs:</label>
      <input type="file" id="xmlFiles" multiple accept=".xml">
    </div>
    <button class="btn" onclick="cargarXMLs()">📤 Cargar XMLs</button>
  </div>
  <div id="xmlResults"></div>
</div>

<!-- TAB: INFORME -->
<div id="informe" class="tab-content">
  <h2>Generar Informe de Citas</h2>
  <div class="card">
    <div class="form-group">
      <label>Modalidad:</label>
      <select id="informeModalidad">
        <option value="UNAM">UNAM (Todos los tipos A+B+C)</option>
        <option value="SNI">SNI/SECIHTI (Solo A+B)</option>
      </select>
    </div>
    <button class="btn" onclick="generarInforme()">📊 Generar Vista Previa</button>
    <button class="btn btn-secondary" onclick="exportarInforme()">💾 Exportar a DOCX</button>
  </div>
  <div id="informePreview"></div>
</div>

<!-- TAB: RIZOMA -->
<div id="rizoma" class="tab-content">
  <h2>Automatización de Rizoma</h2>
  <div class="card">
    <p><strong>⚠️ Importante:</strong> Esta función automatizará el registro de artículos en Rizoma. Asegúrate de haber generado el informe SNI primero.</p>
    <div class="form-group">
      <label>Usuario Rizoma:</label>
      <input type="text" id="rizUser" placeholder="tu@email.com">
    </div>
    <div class="form-group">
      <label>Contraseña:</label>
        <input type="password" id="rizPass" placeholder="••••••••">
    </div>
    <button class="btn" id="rizBtn" onclick="iniciarRizoma()">🤖 Iniciar Automatización</button>
    
    <div id="rizProg" style="display:none;">
      <div class="progress-bar">
        <div class="progress-fill" id="rizProgBar" style="width:0%">
          <span id="rizProgLabel">Iniciando...</span>
        </div>
      </div>
      <div class="status-indicator">
        <div class="dot" id="statusDot"></div>
        <span id="statusText">Listo</span>
      </div>
    </div>
  </div>
  
  <div class="log-console" id="logConsole"></div>
</div>

</div>

<script>
const $ = id => document.getElementById(id);
let activeProfile = null;
let pollInt = null;

function showTab(name) {
  document.querySelectorAll('.tab').forEach(t => t.classList.remove('active'));
  document.querySelectorAll('.tab-content').forEach(c => c.classList.remove('active'));
  event.target.classList.add('active');
  $(name).classList.add('active');
}

async function loadProfiles() {
  const res = await fetch('/api/profiles');
  const data = await res.json();
  const list = $('profilesList');
  if (data.profiles.length === 0) {
    list.innerHTML = '<p>No hay perfiles creados. Crea uno para empezar.</p>';
    return;
  }
  list.innerHTML = data.profiles.map(p => 
    `<div class="profile-item ${p.id === activeProfile ? 'selected' : ''}" onclick="selectProfile('${p.id}')">
      <strong>${p.nombre}</strong><br>
      <small>${p.entidad}</small>
    </div>`
  ).join('');
  if (activeProfile) {
    $('activeName').textContent = data.profiles.find(p => p.id === activeProfile)?.nombre || 'Ninguno';
  }
}

function selectProfile(id) {
  activeProfile = id;
  loadProfiles();
  loadPublicaciones();
  showTab('publicaciones');
}

function showNewProfileForm() {
  $('newProfileForm').style.display = 'block';
}

function hideNewProfileForm() {
  $('newProfileForm').style.display = 'none';
}

async function crearPerfil() {
  const data = {
    nombre: $('profNombre').value,
    entidad: $('profEntidad').value,
    especialista: $('profEspecialista').value,
    revisor: $('profRevisor').value,
    fuentes: $('profFuentes').value.split(',').map(s => s.trim()).filter(Boolean),
    fecha_revision: new Date().toISOString().split('T')[0]
  };
  const res = await fetch('/api/profiles', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify(data)
  });
  const result = await res.json();
  activeProfile = result.id;
  hideNewProfileForm();
  loadProfiles();
  alert('Perfil creado exitosamente');
}

async function loadPublicaciones() {
  if (!activeProfile) return;
  const res = await fetch(`/api/profiles/${activeProfile}`);
  const data = await res.json();
  const pubs = data.publicaciones || [];
  const list = $('pubsList');
  if (pubs.length === 0) {
    list.innerHTML = '<p>No hay publicaciones. Agrega una para comenzar.</p>';
    return;
  }
  list.innerHTML = pubs.map((p, i) => 
    `<div class="pub-item">
      <strong>Publicación ${i + 1}</strong><br>
      ${p.referencia_apa}<br>
      <small>DOI: ${p.doi || 'N/A'} | FI: ${p.fi || 'N/A'} | Cuartil: ${p.quartil || 'N/A'}</small><br>
      <small>Librería: ${p.libreria || 'N/A'}</small><br>
      <small>Citas cargadas: ${p.citas?.length || 0}</small>
    </div>`
  ).join('');
}

function showNewPubForm() {
  if (!activeProfile) { alert('Selecciona un perfil primero'); return; }
  $('newPubForm').style.display = 'block';
}

function hideNewPubForm() {
  $('newPubForm').style.display = 'none';
}

async function agregarPublicacion() {
  if (!activeProfile) return;
  const data = {
    referencia_apa: $('pubRef').value,
    doi: $('pubDoi').value,
    fi: $('pubFi').value,
    quartil: $('pubQuartil').value,
    issn_e: $('pubIssnE').value,
    issn_i: $('pubIssnI').value,
    libreria: $('pubLibreria').value,
    citas: []
  };
  await fetch(`/api/profiles/${activeProfile}/publicaciones`, {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify(data)
  });
  hideNewPubForm();
  loadPublicaciones();
  alert('Publicación agregada');
}

async function cargarXMLs() {
  if (!activeProfile) { alert('Selecciona un perfil primero'); return; }
  const files = $('xmlFiles').files;
  if (files.length === 0) { alert('Selecciona al menos un archivo XML'); return; }
  
  const formData = new FormData();
  formData.append('profile_id', activeProfile);
  for (let f of files) {
    formData.append('files', f);
  }
  
  const res = await fetch('/api/upload_xmls', { method: 'POST', body: formData });
  const data = await res.json();
  
  let html = `<div class="card" style="margin-top:20px;">
    <h3>Resultados de Carga</h3>
    <p>Archivos procesados: ${data.loaded}</p>`;
  
  if (data.files && data.files.length > 0) {
    html += '<h4>Archivos Cargados:</h4><ul>';
    data.files.forEach(f => {
      html += `<li>${f.name}: ${f.citas} citas (A:${f.a}, B:${f.b}, C:${f.c})</li>`;
    });
    html += '</ul>';
  }
  
  if (data.errors && data.errors.length > 0) {
    html += '<h4 style="color:red;">Errores:</h4><ul>';
    data.errors.forEach(e => {
      html += `<li>${e}</li>`;
    });
    html += '</ul>';
  }
  
  html += '</div>';
  $('xmlResults').innerHTML = html;
  loadPublicaciones();
}

async function generarInforme() {
  if (!activeProfile) { alert('Selecciona un perfil primero'); return; }
  const modalidad = $('informeModalidad').value;
  const res = await fetch(`/api/profiles/${activeProfile}/informe?modalidad=${modalidad}`);
  const data = await res.json();
  
  let html = `<div class="card" style="margin-top:20px;">
    <h3>Informe de Citas - ${data.modalidad}</h3>
    <p><strong>Investigador:</strong> ${data.investigador}</p>
    <p><strong>Entidad:</strong> ${data.entidad}</p>
    <p><strong>Fecha:</strong> ${data.fecha_revision}</p>
    <hr>
    <h4>Resumen</h4>
    <p>Total Tipo A (Individual): ${data.total_a}</p>
    <p>Total Tipo B (Colaborativa): ${data.total_b}</p>`;
  
  if (data.modalidad === 'UNAM') {
    html += `<p>Total Tipo C (Autocita): ${data.total_c}</p>`;
  }
  
  html += `<p><strong>TOTAL GENERAL: ${data.total_general}</strong></p>
    <hr>
    <h4>Documentos Procesados</h4>`;
  
  data.documentos.forEach((doc, i) => {
    html += `<div class="pub-item">
      <strong>Documento ${i + 1}</strong><br>
      ${doc.referencia}<br>
      <small>Citas: A=${doc.cnt_a}, B=${doc.cnt_b}`;
    if (data.modalidad === 'UNAM') {
      html += `, C=${doc.cnt_c}`;
    }
    html += `</small>
    </div>`;
  });
  
  html += '</div>';
  $('informePreview').innerHTML = html;
}

async function exportarInforme() {
  if (!activeProfile) { alert('Selecciona un perfil primero'); return; }
  const modalidad = $('informeModalidad').value;
  window.location.href = `/api/profiles/${activeProfile}/export_docx?modalidad=${modalidad}`;
}

function log(msg, cls = '') {
  const console = $('logConsole');
  const line = document.createElement('div');
  line.className = `log-line ${cls}`;
  line.textContent = msg;
  console.appendChild(line);
  console.scrollTop = console.scrollHeight;
}

async function iniciarRizoma() {
  if (!activeProfile) { alert('Selecciona un perfil primero'); return; }
  const creds = { usuario: $('rizUser').value, password: $('rizPass').value };
  if (!creds.usuario || !creds.password) { alert('Ingresa usuario y contraseña de Rizoma'); return; }
  
  $('rizBtn').disabled = true;
  $('rizProg').style.display = 'block';
  $('statusDot').className = 'dot busy';
  $('statusText').textContent = 'Automatizando Rizoma...';
  log('🌐 Iniciando automatización en Rizoma (SNI)...', 'log-info');

  await fetch(`/api/profiles/${activeProfile}/rizoma`, {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({ credenciales: creds, modalidad: 'SNI' })
  });

  pollInt = setInterval(async () => {
    const res = await fetch('/api/task_status');
    const s = await res.json();
    s.log.forEach(l => log(l, l.startsWith('✅') || l.startsWith('✓') ? 'log-ok' : l.startsWith('⚠') || l.startsWith('Error') ? 'log-err' : ''));
    if (s.total > 0) {
      $('rizProgBar').style.width = Math.round(s.progress / s.total * 100) + '%';
      $('rizProgLabel').textContent = `Artículo ${s.progress} de ${s.total}`;
    }
    if (s.status === 'done' || s.status === 'error') {
      clearInterval(pollInt);
      $('rizBtn').disabled = false;
      $('statusDot').className = 'dot';
      $('statusText').textContent = 'Listo';
      log(s.status === 'done' ? '✅ Registro en Rizoma completado' : '⚠ Proceso con errores', s.status === 'done' ? 'log-ok' : 'log-err');
      await fetch('/api/task_reset', { method: 'POST' });
    }
  }, 800);
}

loadProfiles();
</script>
</body>
</html>"""

# ══════════════════════════════════════════════════════════════════════════════
# RUTAS API
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    return render_template_string(HTML)

@app.route("/api/profiles")
def list_profiles():
    profiles = load_profiles()
    return jsonify({"profiles": [{"id": k, "nombre": v["nombre"], "entidad": v["entidad"]} for k, v in profiles.items()]})

@app.route("/api/profiles", methods=["POST"])
def create_profile():
    data = request.json
    profiles = load_profiles()
    pid = re.sub(r'\W+', '_', data["nombre"].lower())[:30] + f"_{len(profiles)}"
    profiles[pid] = {**data, "publicaciones": data.get("publicaciones", [])}
    save_profiles(profiles)
    return jsonify({"id": pid})

@app.route("/api/profiles/<pid>")
def get_profile_route(pid):
    p = get_profile(pid)
    if not p:
        return jsonify({"error": "Not found"}), 404
    return jsonify({**p, "id": pid})

@app.route("/api/profiles/<pid>", methods=["DELETE"])
def delete_profile(pid):
    profiles = load_profiles()
    profiles.pop(pid, None)
    save_profiles(profiles)
    return jsonify({"ok": True})

@app.route("/api/profiles/<pid>/publicaciones", methods=["POST"])
def add_pub(pid):
    profiles = load_profiles()
    if pid not in profiles:
        return jsonify({"error": "Not found"}), 404
    profiles[pid]["publicaciones"].append(request.json)
    save_profiles(profiles)
    return jsonify({"ok": True})

@app.route("/api/profiles/<pid>/publicaciones/<int:idx>", methods=["DELETE"])
def del_pub(pid, idx):
    profiles = load_profiles()
    pubs = profiles[pid].get("publicaciones", [])
    if 0 <= idx < len(pubs):
        pubs.pop(idx)
    profiles[pid]["publicaciones"] = pubs
    save_profiles(profiles)
    return jsonify({"ok": True})

@app.route("/api/upload_xmls", methods=["POST"])
def upload_xmls():
    pid = request.form.get("profile_id")
    profiles = load_profiles()
    if pid not in profiles:
        return jsonify({"error": "Perfil no encontrado"}), 404

    files_info = []
    errors = []
    loaded = 0

    for file in request.files.getlist("files"):
        if not file.filename.endswith(".xml"):
            errors.append(f"{file.filename}: no es un XML válido")
            continue

        save_path = UPLOAD_DIR / file.filename
        file.save(str(save_path))

        citas = parse_endnote_xml(str(save_path))
        stem = Path(file.filename).stem

        cnt_a = sum(1 for c in citas if c["tipo"] == "A")
        cnt_b = sum(1 for c in citas if c["tipo"] == "B")
        cnt_c = sum(1 for c in citas if c["tipo"] == "C")

        files_info.append({"name": file.filename, "citas": len(citas), "a": cnt_a, "b": cnt_b, "c": cnt_c})

        for pub in profiles[pid]["publicaciones"]:
            if pub.get("libreria", "").strip() == stem:
                pub["citas"] = citas
                break
        else:
            errors.append(f"{file.filename}: no se encontró publicación con librería '{stem}'")

        loaded += 1

    save_profiles(profiles)
    return jsonify({"loaded": loaded, "files": files_info, "errors": errors})

@app.route("/api/profiles/<pid>/informe")
def get_informe(pid):
    profile = get_profile(pid)
    if not profile:
        return jsonify({"error": "Not found"}), 404
    modalidad = request.args.get("modalidad", "UNAM")
    return jsonify(generar_informe(profile, modalidad))

@app.route("/api/profiles/<pid>/export_docx")
def export_docx(pid):
    profile = get_profile(pid)
    if not profile:
        return jsonify({"error": "Not found"}), 404
    modalidad = request.args.get("modalidad", "UNAM")
    informe = generar_informe(profile, modalidad)

    nombre_safe = re.sub(r'\W+', '_', profile.get("nombre", "informe"))
    out_name = f"Informe_Citas_{nombre_safe}_{modalidad}_{datetime.now().strftime('%Y%m%d')}.docx"
    out_path = str(OUTPUT_DIR / out_name)

    if DOCX_OK:
        exportar_docx(informe, out_path)
        return send_from_directory(str(OUTPUT_DIR), out_name, as_attachment=True)
    else:
        return jsonify({"error": "python-docx no está instalado. Instala con: pip install python-docx"}), 500

@app.route("/api/profiles/<pid>/rizoma", methods=["POST"])
def start_rizoma(pid):
    profile = get_profile(pid)
    if not profile:
        return jsonify({"error": "Not found"}), 404

    body = request.json
    credenciales = body.get("credenciales", {})
    modalidad = body.get("modalidad", "SNI")
    informe = generar_informe(profile, modalidad)

    _task.update({"active": True, "status": "running", "log": [], "progress": 0, "total": 0})

    def run():
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            loop.run_until_complete(automatizar_rizoma(profile, informe, credenciales))
        except Exception as e:
            _task["log"].append(f"Error: {e}")
            _task["status"] = "error"
        finally:
            loop.close()

    t = threading.Thread(target=run, daemon=True)
    t.start()
    return jsonify({"ok": True})

@app.route("/api/task_status")
def task_status():
    logs = _task["log"][:]
    _task["log"] = []
    return jsonify({**_task, "log": logs})

@app.route("/api/task_reset", methods=["POST"])
def task_reset():
    _task.update({"active": False, "status": "idle", "log": [], "progress": 0, "total": 0})
    return jsonify({"ok": True})

# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("=" * 60)
    print("  Sistema de Automatización Rizoma — BCCT/UNAM")
    print("  Abre tu navegador en: http://localhost:5000")
    print("=" * 60)
    print()
    print("📦 Dependencias instaladas:")
    print(f"  - python-docx: {'✅' if DOCX_OK else '❌ (pip install python-docx)'}")
    print(f"  - reportlab: {'✅' if PDF_OK else '❌ (pip install reportlab)'}")
    print(f"  - playwright: {'✅' if PLAYWRIGHT_OK else '❌ (pip install playwright)'}")
    print()
    print("=" * 60)
    app.run(debug=False, port=5000, threaded=True)